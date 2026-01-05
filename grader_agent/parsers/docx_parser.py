"""Simple parser for Word document files."""

import re
import logging
from typing import Tuple, Dict, Optional
from docx import Document

logger = logging.getLogger(__name__)


class DocxParser:
    """Parse Word documents to extract text content."""
    
    def __init__(self, path: str):
        """Initialize parser with document path."""
        self.path = path
        self.document = None
    
    def parse(self, use_llm_extraction: bool = True, api_key: Optional[str] = None) -> Tuple[str, str, Dict[str, str]]:
        """
        Parse the document and extract content.
        
        Returns:
            Tuple of (student_id, student_name, answers_dict)
        """
        try:
            self.document = Document(self.path)
        except Exception as e:
            logger.error(f"Failed to read document {self.path}: {e}")
            raise
        
        student_id = "unknown"
        student_name = None
        answers = {}
        
        # First, try to extract student info from the first few paragraphs
        # Student info is usually at the top of the document
        logger.info(f"Parsing document: {self.path}")
        
        # Check first 30 paragraphs for student info (usually at the top)
        early_text = []
        for i, para in enumerate(self.document.paragraphs[:30]):
            text = para.text.strip()
            if text:  # Only add non-empty paragraphs
                early_text.append(text)
                logger.debug(f"Paragraph {i}: {text[:100]}")  # Log first 100 chars
        
        early_text_str = "\n".join(early_text)
        logger.info(f"Early text (first 30 paragraphs):\n{early_text_str[:500]}")  # Log first 500 chars
        
        # Multiple patterns to try for NetID/Student ID
        netid_patterns = [
            r"NetID\s*[:=]\s*(\S+)",  # NetID: abc123 or NetID=abc123
            r"Net\s*ID\s*[:=]\s*(\S+)",  # Net ID: abc123
            r"Student\s*ID\s*[:=]\s*(\S+)",  # Student ID: abc123
            r"ID\s*[:=]\s*(\S+)",  # ID: abc123 (more generic)
            r"(\w{2,3}\d{5,})",  # Pattern like abc12345 (2-3 letters + 5+ digits)
        ]
        
        # Try each pattern
        for pattern in netid_patterns:
            match = re.search(pattern, early_text_str, re.IGNORECASE)
            if match:
                student_id = match.group(1).strip()
                logger.info(f"Found student_id using pattern '{pattern}': {student_id}")
                break
        
        # Multiple patterns for student name/author
        name_patterns = [
            r"Author\s*[:=]\s*(.+?)(?:\n|$)",  # Author: Name (until newline or end)
            r"Name\s*[:=]\s*(.+?)(?:\n|$)",  # Name: Student Name
            r"Student\s*Name\s*[:=]\s*(.+?)(?:\n|$)",  # Student Name: ...
            r"Submitted\s+by\s*[:=]\s*(.+?)(?:\n|$)",  # Submitted by: ...
        ]
        
        # Try each pattern
        for pattern in name_patterns:
            match = re.search(pattern, early_text_str, re.IGNORECASE | re.MULTILINE)
            if match:
                student_name = match.group(1).strip()
                # Clean up - remove extra whitespace and common suffixes
                student_name = re.sub(r'\s+', ' ', student_name)  # Normalize whitespace
                student_name = re.sub(r'\s*[\(\[].*?[\)\]]\s*', '', student_name)  # Remove anything in parentheses/brackets
                student_name = student_name.strip()
                if student_name and len(student_name) > 2:  # Valid name should be at least 3 chars
                    logger.info(f"Found student_name using pattern '{pattern}': {student_name}")
                    break
        
        # If still not found, check tables (sometimes student info is in a table)
        if student_id == "unknown" or student_name is None:
            logger.info("Checking tables for student info...")
            for table in self.document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        # Check for NetID in table
                        for pattern in netid_patterns:
                            match = re.search(pattern, row_text, re.IGNORECASE)
                            if match and student_id == "unknown":
                                student_id = match.group(1).strip()
                                logger.info(f"Found student_id in table: {student_id}")
                        
                        # Check for name in table
                        for pattern in name_patterns:
                            match = re.search(pattern, row_text, re.IGNORECASE)
                            if match and student_name is None:
                                student_name = match.group(1).strip()
                                student_name = re.sub(r'\s+', ' ', student_name)
                                student_name = re.sub(r'\s*[\(\[].*?[\)\]]\s*', '', student_name)
                                student_name = student_name.strip()
                                if student_name and len(student_name) > 2:
                                    logger.info(f"Found student_name in table: {student_name}")
                                    break
        
        # Collect all text for answers extraction
        all_text_parts = []
        for para in self.document.paragraphs:
            all_text_parts.append(para.text)
        
        for table in self.document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                all_text_parts.append(row_text)
        
        full_text = "\n".join(all_text_parts)
        
        # Try LLM extraction if enabled and regex didn't find both
        if use_llm_extraction and (student_id == "unknown" or student_name is None):
            try:
                from .student_info_extractor import StudentInfoExtractor
                extractor = StudentInfoExtractor(api_key=api_key)
                llm_id, llm_name = extractor.extract_from_text(full_text, self.path)
                
                # Use LLM results if regex didn't find them
                if student_id == "unknown" and llm_id != "unknown":
                    student_id = llm_id
                    logger.info(f"LLM found student_id: {student_id}")
                
                if student_name is None and llm_name:
                    student_name = llm_name
                    logger.info(f"LLM found student_name: {student_name}")
            except Exception as e:
                logger.warning(f"LLM extraction failed, using regex results: {e}")
        
        # Log final extracted info
        logger.info(f"Final extraction - student_id: {student_id}, student_name: {student_name}")
        
        # Try to find questions
        question_pattern = re.compile(
            r"^Q(?P<qid>[0-9a-zA-Z._-]+)|(?:Question|Problem)\s*(?P<qid2>\d+(?:\.\w+)?)",
            re.IGNORECASE | re.MULTILINE
        )
        
        question_matches = list(question_pattern.finditer(full_text))
        
        if question_matches:
            for i, match in enumerate(question_matches):
                qid = match.group("qid") or match.group("qid2")
                if not qid:
                    continue
                
                start_pos = match.end()
                end_pos = question_matches[i + 1].start() if i + 1 < len(question_matches) else len(full_text)
                answer_text = full_text[start_pos:end_pos].strip()
                
                if answer_text:
                    answers[qid] = answer_text
        else:
            # No questions found, return all text
            answers["all"] = full_text
        
        return student_id, student_name, answers
